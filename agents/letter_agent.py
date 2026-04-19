import os
import json
from groq import Groq
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import locale
locale.setlocale(locale.LC_TIME, "fr_FR.UTF-8")

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))


def generate_letter_text(cv: dict, job: dict) -> str:
    """
    Demande à Groq de rédiger la lettre de motivation.
    Retourne le texte brut de la lettre.
    """

    analysis = job.get("analysis", {})

    prompt = f"""
Tu es un expert en recrutement tech. Rédige une lettre de motivation professionnelle et personnalisée.

Profil du candidat :
- Nom        : {cv['name']}
- Titre      : {cv['title']}
- Résumé     : {cv['summary']}
- Skills AI  : {', '.join(cv['skills']['ai'])}
- Skills Dev : {', '.join(cv['skills']['backend'] + cv['skills']['frontend'])}
- Expérience principale : {cv['experiences'][0]['role']} chez {cv['experiences'][0]['company']}
- Projet phare : {cv['projects'][0]['name']} — {cv['projects'][0]['description']}

Offre ciblée :
- Poste      : {job['title']}
- Entreprise : {job['company']}
- Missions   : {', '.join(analysis.get('key_missions', []))}
- Skills requis : {', '.join(analysis.get('required_skills', []))}
- Ton de la boîte : {analysis.get('company_tone', 'professionnel')}
- Mots-clés à utiliser : {', '.join(analysis.get('match_keywords', []))}

Règles de rédaction :
- Longueur : 3 paragraphes maximum, concis et impactants
- Ton : adapte-toi au company_tone (startup = dynamique, corporate = formel)
- Paragraphe 1 : accroche + pourquoi ce poste m'intéresse
- Paragraphe 2 : 2 réalisations concrètes avec chiffres qui matchent l'offre
- Paragraphe 3 : conclusion + disponibilité (juillet 2026)
- NE PAS inclure : objet, date, formule de politesse finale — juste les 3 paragraphes
- Langue : français
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,  # ← plus élevé qu'analyzer : on veut du texte naturel
    )

    return response.choices[0].message.content.strip()


def generate_letter_docx(cv: dict, job: dict, output_dir: str = "output") -> str:
    """
    Génère la lettre de motivation en DOCX.
    """

    os.makedirs(output_dir, exist_ok=True)

    # Génération du texte
    print("Rédaction de la lettre en cours...")
    letter_text = generate_letter_text(cv, job)
    print("✓ Texte généré\n")

    doc = Document()

    # Style global
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Header candidat ---
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(cv["name"])
    name_run.bold = True
    name_run.font.size = Pt(14)

    contact_p = doc.add_paragraph()
    contact_p.add_run(
        f"{cv['email']}  |  {cv['phone']}  |  {cv['linkedin']}"
    ).font.size = Pt(10)

    doc.add_paragraph()  # espace

    # --- Date ---
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f"Casablanca, le {date.today().strftime('%d %B %Y')}")

    doc.add_paragraph()

    # --- Destinataire ---
    doc.add_paragraph(f"À l'attention du service recrutement")
    dest_p = doc.add_paragraph()
    dest_p.add_run(job["company"]).bold = True

    doc.add_paragraph()

    # --- Objet ---
    obj_p = doc.add_paragraph()
    obj_run = obj_p.add_run(f"Objet : Candidature au poste de {job['title']}")
    obj_run.bold = True
    obj_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph()

    # --- Formule d'appel ---
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()

    # --- Corps de la lettre (3 paragraphes max) ---
    paragraphs = [p.strip() for p in letter_text.split("\n") if p.strip()]
    for paragraph in paragraphs[:3]:
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # --- Signature ---
    doc.add_paragraph("Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.")
    doc.add_paragraph()
    signature_p = doc.add_paragraph()
    signature_p.add_run(cv["name"]).bold = True

    # --- Sauvegarde ---
    company_name = job["company"].replace(" ", "_").replace("/", "-")
    filename = f"{output_dir}/LM_Oussama_{company_name}.docx"
    doc.save(filename)
    print(f"✓ Lettre générée : {filename}")
    return filename


def run_letter_agent(job: dict) -> str:
    """
    Pipeline complet pour une offre :
    1. Charger le CV master
    2. Générer le contenu de la lettre
    3. Sauvegarder en DOCX
    """
    with open("data/cv_master.json", "r", encoding="utf-8") as f:
        cv = json.load(f)

    return generate_letter_docx(cv, job)


if __name__ == "__main__":
    # On prend la première offre analysée pour tester
    with open("data/jobs_analyzed.json", "r", encoding="utf-8") as f:
        jobs = json.load(f)

    first_job = jobs[0]
    print(f"Génération de la lettre pour : {first_job['title']} — {first_job['company']}\n")

    filepath = run_letter_agent(first_job)
    print(f"\nFichier créé : {filepath}")