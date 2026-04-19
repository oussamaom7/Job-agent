import os
import json
from groq import Groq
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))


def load_cv_master(filepath: str = "data/cv_master.json") -> dict:
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def select_content(cv: dict, job_analysis: dict) -> dict:
    """
    Demande à Groq de choisir quels blocs du CV mettre en avant
    selon l'analyse de l'offre.
    Retourne un dict avec les IDs sélectionnés et les reformulations.
    """

    prompt = f"""
Tu es un expert en recrutement tech. Tu dois personnaliser un CV pour une offre d'emploi.

Analyse de l'offre :
{json.dumps(job_analysis, ensure_ascii=False, indent=2)}

Contenu disponible dans le CV :
- Expériences : {[e['id'] for e in cv['experiences']]}
- Projets     : {[p['id'] for p in cv['projects']]}
- Skills      : {list(cv['skills'].keys())}

Retourne UNIQUEMENT un JSON valide avec ce format :
{{
  "selected_experiences": ["id1", "id2"],
  "selected_projects"   : ["id1", "id2"],
  "selected_skills"     : ["frontend", "backend", "ai", "devops", "languages"],
  "cv_title"            : "titre du poste adapté à l'offre",
  "summary"             : "résumé personnalisé de 2 phrases maximum pour cette offre"
}}

Règles :
- Sélectionne max 2 expériences et max 3 projets
- Priorise les éléments dont les tags matchent les required_skills de l'offre
- Le cv_title doit reprendre les mots-clés exacts de l'offre
- Le summary doit mentionner 2-3 compétences clés de l'offre
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )

    raw = response.choices[0].message.content
    
    try:
        # On nettoie les balises markdown si Groq en ajoute
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            # Supprime la première ligne (```json) et la dernière (```)
            cleaned = "\n".join(cleaned.split("\n")[1:-1])
        return json.loads(cleaned)
    except json.JSONDecodeError:
        print(f"⚠ JSON invalide depuis Groq")
        print(f"Reçu : {raw[:200]}")  # ← on affiche ce qu'on a reçu pour débugger
        return {}


def generate_docx(cv: dict, selection: dict, job: dict, output_dir: str = "output") -> str:
    """
    Génère un fichier DOCX propre à partir du CV master et de la sélection.
    Retourne le chemin du fichier créé.
    """

    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # --- Styles globaux ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        return p

    def add_separator():
        p = doc.add_paragraph()
        run = p.add_run("─" * 60)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(9)

    # --- Header ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv["name"])
    name_run.bold = True
    name_run.font.size = Pt(20)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(selection.get("cv_title", cv["title"]))

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run(
        f"{cv['email']}  |  {cv['phone']}  |  {cv['linkedin']}  |  {cv['github']}"
    ).font.size = Pt(9)

    add_separator()

    # --- Summary ---
    add_heading("Summary", level=2)
    doc.add_paragraph(selection.get("summary", cv["summary"]))

    # --- Skills ---
    add_heading("Technical Skills", level=2)
    add_separator()
    selected_skill_cats = selection.get("selected_skills", list(cv["skills"].keys()))
    
    for category in selected_skill_cats:
        if category in cv["skills"]:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{category.capitalize()}: ")
            run_label.bold = True
            p.add_run(", ".join(cv["skills"][category]))

    # --- Experience ---
    add_heading("Experience", level=2)
    add_separator()
    selected_exp_ids = selection.get("selected_experiences", [e["id"] for e in cv["experiences"]])
    
    for exp in cv["experiences"]:
        if exp["id"] in selected_exp_ids:
            p = doc.add_paragraph()
            r = p.add_run(f"{exp['role']}  —  {exp['company']}")
            r.bold = True
            doc.add_paragraph(exp["period"]).runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            for bullet in exp["bullets"]:
                doc.add_paragraph(bullet, style="List Bullet")

    # --- Projects ---
    add_heading("Projects", level=2)
    add_separator()
    selected_proj_ids = selection.get("selected_projects", [p["id"] for p in cv["projects"]])
    
    for proj in cv["projects"]:
        if proj["id"] in selected_proj_ids:
            p = doc.add_paragraph()
            p.add_run(proj["name"]).bold = True
            doc.add_paragraph(proj["description"])
            for highlight in proj["highlights"]:
                doc.add_paragraph(highlight, style="List Bullet")

    # --- Education ---
    add_heading("Education", level=2)
    add_separator()
    edu = cv["education"]
    p = doc.add_paragraph()
    p.add_run(f"{edu['degree']}  —  {edu['school']}").bold = True
    doc.add_paragraph(f"{edu['period']}  |  {edu['note']}")

    # --- Sauvegarde ---
    company_name = job["company"].replace(" ", "_").replace("/", "-")
    filename = f"{output_dir}/CV_Oussama_{company_name}.docx"
    doc.save(filename)
    print(f"✓ CV généré : {filename}")
    return filename


def run_cv_agent(job: dict) -> str:
    """
    Pipeline complet pour un job :
    1. Charge le CV master
    2. Sélectionne le contenu via Groq
    3. Génère le DOCX
    """
    cv        = load_cv_master()
    analysis  = job.get("analysis", {})
    selection = select_content(cv, analysis)
    filepath  = generate_docx(cv, selection, job)
    return filepath


if __name__ == "__main__":
    # On charge la première offre analysée pour tester
    with open("data/jobs_analyzed.json", "r", encoding="utf-8") as f:
        jobs = json.load(f)

    first_job = jobs[0]
    print(f"Génération du CV pour : {first_job['title']} — {first_job['company']}\n")
    
    filepath = run_cv_agent(first_job)
    print(f"\nFichier créé : {filepath}")